# from django.http import HttpResponse, JsonResponse
# from docx import Document
# from docx.enum.table import WD_ALIGN_VERTICAL
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.shared import Pt, Inches
#
# from myapp.models import User, OwnSkill
#
#
# def create_resume(request):
#
#     id=request.POST['lid']
#     vid=request.POST['vid']
#     u=User.objects.get(LOGIN=id)
#     s=OwnSkill.objects.filter(USER__LOGIN_id=id)
#     print(s)
#     document = Document()
#     import datetime
#
#
#
#     # Add a title and center it
#     title = document.add_heading(u.name, level=1)
#     title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#     # Create a two-column section
#     section = document.sections[-1]
#     section.start_type
#     section.left_margin = Inches(1)
#     section.right_margin = Inches(1)
#
#     # Add a photo
#     photo_path = 'C:\\Users\\Lenovo\\PycharmProjects\\career_path_navigator\\'+u.photo  # Replace with the actual path to your photo
#     document.add_picture(photo_path, width=Inches(2.0), height=Inches(2.0))
#
#     # Personal Information Section
#     add_info_section(document, 'Address', u.place +','+ u.pin +','+ u.district)
#     add_info_section(document, 'Phone', u.phone)
#     add_info_section(document, 'Email', u.email)
#
#     # Create a two-column table for the rest of the content
#     content_table = document.add_table(rows=1, cols=2)
#     content_table.alignment = WD_ALIGN_VERTICAL.CENTER
#     content_table.allow_autofit = False
#     content_table.autofit = False
#     content_table.columns[0].width = Inches(2.5)
#     content_table.columns[1].width = Inches(3.5)
#
#     # Education Section
#     add_education_section(content_table,"Qualification"+ u.qualif)
#
#     # Work Experience Section
#     add_experience_section(content_table,"Experience" + u.exp)
#
#     # Skills Section
#     l = ', '.join(i.SKILL.skill for i in s)
#     print(str(l))
#     add_skills_section(content_table,l)
#
#         # Save the document
#
#     date ='/static/cv/' + datetime.datetime.now().strftime('%Y%m%d-%H%M%S')+u.name+'.docx'
#     document.save('C:/Users/Lenovo/PycharmProjects/career_path_navigator/myapp'+date)
#
#
#     da = datetime.datetime.now().date()
#     if not JobRequest.objects.filter(USER=u,VACCANCY_id=vid).exists():
#         j=JobRequest()
#         j.date=da
#         j.USER=User.objects.get(LOGIN=id)
#         j.VACCANCY_id=vid
#         j.status='pending'
#         j.file=date
#         j.save()
#         return JsonResponse({"status":"ok"})
#     return JsonResponse({"status":"no"})
#
#
# def add_info_section(document, label, content):
#     paragraph = document.add_paragraph()
#     paragraph.add_run(label + ':').bold = True
#     paragraph.add_run(' ' + content)
#
# def add_education_section(content_table, degree):
#     row = content_table.add_row().cells
#     cell = row[0]
#     paragraph = cell.add_paragraph()
#     paragraph.add_run(degree).bold = True
#     paragraph.add_run()
#
# def add_experience_section(content_table, title):
#     row = content_table.add_row().cells
#     cell = row[0]
#     paragraph = cell.add_paragraph()
#     paragraph.add_run(title).bold = True
#
# def add_skills_section(content_table, *skills):
#     row = content_table.add_row().cells
#     cell = row[0]
#     paragraph = cell.add_paragraph()
#     paragraph.add_run('Skills').bold = True
#     paragraph.add_run('\n' + ', '.join(skills))
#
from django.http import HttpResponse, JsonResponse
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
import datetime

from myapp.models import User, OwnSkill, JobRequest

def create_resume(request):
    id = request.POST['lid']
    vid = request.POST['vid']
    u = User.objects.get(LOGIN=id)
    s = OwnSkill.objects.filter(USER__LOGIN_id=id)

    document = Document()

    # Add a title and center it
    title = document.add_heading(u.name, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set document margins
    section = document.sections[-1]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add a photo
    photo_path = 'C:\\Users\\Lenovo\\PycharmProjects\\career_path_navigator\\' + u.photo  # Replace with actual path
    document.add_picture(photo_path, width=Inches(2.0), height=Inches(2.0))

    # Personal Information Section
    add_info_section(document, 'Address', u.place + ', ' + u.pin + ', ' + u.district)
    add_info_section(document, 'Phone', str(u.phone))  # Fixed issue: Convert phone number to string
    add_info_section(document, 'Email', u.email)

    # Create a table for other sections
    content_table = document.add_table(rows=1, cols=2)
    content_table.alignment = WD_ALIGN_VERTICAL.CENTER
    content_table.allow_autofit = False

    # Education Section
    add_education_section(content_table, "Qualification: " + u.qualification)

    # Work Experience Section
    add_experience_section(content_table, "Experience: " + u.experience)

    # Skills Section
    skills_list = ', '.join(i.SKILL.skill for i in s)
    add_skills_section(content_table, skills_list)

    # Save the document
    date = '/static/cv/' + datetime.datetime.now().strftime('%Y%m%d-%H%M%S') + u.name + '.docx'
    document.save('C:/Users/Lenovo/PycharmProjects/career_path_navigator/myapp' + date)

    # Save job request if not exists
    da = datetime.datetime.now().date()
    if not JobRequest.objects.filter(USER=u, VACANCY_id=vid).exists():
        j = JobRequest()
        j.date = da
        j.USER = u
        j.VACANCY_id = vid
        j.status = 'pending'
        j.file = date
        j.save()
        return JsonResponse({"status": "ok"})

    return JsonResponse({"status": "no"})

# Helper functions
def add_info_section(document, label, content):
    paragraph = document.add_paragraph()
    paragraph.add_run(label + ': ').bold = True
    paragraph.add_run(str(content))  # Ensure content is a string

def add_education_section(content_table, degree):
    row = content_table.add_row().cells
    paragraph = row[0].add_paragraph()
    paragraph.add_run(degree).bold = True

def add_experience_section(content_table, title):
    row = content_table.add_row().cells
    paragraph = row[0].add_paragraph()
    paragraph.add_run(title).bold = True

def add_skills_section(content_table, skills):
    row = content_table.add_row().cells
    paragraph = row[0].add_paragraph()
    paragraph.add_run('Skills: ').bold = True
    paragraph.add_run(skills)
