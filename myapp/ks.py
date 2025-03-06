from django.http import JsonResponse
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
import os
# C:\\Users\\91815\\PycharmProjects\\profnet\\media\\20231025-094059.jpg
from docx.shared import Cm, Inches
from win32com import client
import pythoncom

from myapp.models import OwnSkill, User,JobRequest

pythoncom.CoInitialize()

from PIL import Image, ImageTk
print("ok")
# from DBConnection import *
def create_resume3(request):

    id=request.POST['lid']
    vid=request.POST['vid']
    u=User.objects.get(LOGIN=id)
    s=OwnSkill.objects.filter(USER__LOGIN_id=id)
    print(s)
    document = Document()
    # df = pd.read_csv('objective.csv')
    pythoncom.CoInitialize()

    document.add_heading('Resume', 0)
    error = None
    qualifications = ['Btech', 'Mtech', 'BBA']
    sectors = ['IT', 'Mechanical', 'Automobile', 'Electrical']
    experience_level = ['Entry', 'Experienced']
    year_passing = ['2010', '2011', '2012', '2013', '2014', '2015', '2016', '2017', '2018']
    experience = ['0 - 1 yrs', '2 - 4 yrs', '5 - 7 yrs', '8 - 10 yrs', '10 - 12 yrs']

    lid="1"

    photo_path = 'C:\\Users\\Lenovo\\PycharmProjects\\career_path_navigator\\'+str(u.photo)
    document.add_picture(photo_path, width=Inches(2.0), height=Inches(2.0))

    p = document.add_paragraph("")
    p.alignment = 0
    # paragraph = document.add_paragraph("RESUME")
    #
    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p = document.add_picture('C://Users//DELL//PycharmProjects//cv generater//'+res["photo"])
    # last_paragraph = document.paragraphs[-1]
    # last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = document.add_table(rows=2, cols=3)
    # table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells

    # table.rows[0].height = Cm(0.9)

    hdr_cells[0].text = "\n"+ "\n"+ str(u.name)
    hdr_cells[1].text=""
    hdr_cells1 = table.rows[1].cells
    hdr_cells1[0].text = str(u.phone) + str(u.email)
    hdr_cells1[1].text = ''
    hdr_cells1[2].text = ''

    d=hdr_cells[2]
    e=hdr_cells1[2]
    d.merge(e)



    error = None
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p = document.add_paragraph('')
    p.alignment = 0
    p = "aaaa"
    # my_objective ="aaa" #df[(df['Level'] == my_experience_level) & (df['Career'] == my_industry)]['Message'][p]
    p = document.add_paragraph('Address                                                : ' + "")
    p.alignment = 0
    p = document.add_paragraph('                                                                  ' + str(u.place))
    p.alignment = 0
    p = document.add_paragraph('                                                                  ' + "")
    p.alignment = 0
    p = document.add_paragraph('                                                                  ' + str(u.pin))
    p.alignment = 0
    p = document.add_paragraph('                                                                  ' + str(u.district))
    p.alignment = 0
    p = document.add_paragraph(
        'Contact                                                  : ' + str(u.phone))
    p = document.add_paragraph("")
    p = document.add_paragraph("")
    p.alignment = 0
    p = document.add_paragraph('Qualification                                      : ' + str(u.qualification))
    p.alignment = 0
    p = document.add_paragraph('Experience                                         : ' + str(u.experience))
    p.alignment = 0
    l = ', '.join(i.SKILL.skill for i in s)
    p = document.add_paragraph('Skill set                                               : ' + l)
    p.alignment = 0
    p = document.add_paragraph('                                       ' + "")
    p.alignment = 0

    p = document.add_paragraph("")
    pot="I hereby declare that all the information given above is true and correct to the best of my knowledge."
    p = document.add_paragraph(pot)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p = document.add_paragraph("")
    p = document.add_paragraph("Name      : " + str(u.name))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    import datetime
    p = document.add_paragraph("Date  : "+  str(datetime.datetime.now().date()))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date = '/static/cv/' + datetime.datetime.now().strftime('%Y%m%d-%H%M%S') + str(u.name) + '.docx'
    document.save('C:/Users/Lenovo/PycharmProjects/career_path_navigator/myapp' + date)

    da = datetime.datetime.now().date()
    if not JobRequest.objects.filter(USER=u, VACANCY_id=vid).exists():
        j = JobRequest()
        j.date = da
        j.USER = User.objects.get(LOGIN=id)
        j.VACANCY_id = vid
        j.status = 'pending'
        j.file = date
        j.save()
        return JsonResponse({"status": "ok"})
    return JsonResponse({"status": "no"})

    wdFormatPDF = 17

    in_file = os.path.abspath('C:\\Users\\Lenovo\\PycharmProjects\\career_path_navigator\\myapp\\static\\cv\\resume_model1.docx')
    out_file = os.path.abspath('C:\\Users\\Lenovo\\PycharmProjects\\career_path_navigator\\myapp\\static\\cv\\resume_model1.pdv')

    word = client.DispatchEx("Word.Application")
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    # word.Quit()



