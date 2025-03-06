# import datetime
#
# from django.http import JsonResponse
# from docx import Document
# from docx.shared import Inches
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.table import WD_ALIGN_VERTICAL
#
# from myapp.models import User, OwnSkill
#
#
# def create_resume2(request):
#
#     id=request.POST['lid']
#     vid=request.POST['vid']
#     u=User.objects.get(LOGIN=id)
#     s=OwnSkill.objects.filter(USER__LOGIN_id=id)
#     print(s)
#     document = Document()
#
#     # Add a title
#     document.add_heading('Resume', 0)
#
#     # Add a photo
#     photo_path = 'C:\\Users\\Lenovo\\PycharmProjects\\career_path_navigator\\'+u.photo  # Replace with the actual path to your photo
#     document.add_picture(photo_path, width=Inches(2.0), height=Inches(2.0))
#
#     # Add personal information
#     personal_info = document.add_table(rows=3, cols=2)
#     personal_info.alignment = WD_ALIGN_VERTICAL.CENTER
#
#     personal_info.rows[0].cells[0].text = 'Name:'
#     personal_info.rows[0].cells[1].text = u.name
#     personal_info.rows[1].cells[0].text = 'Address:'
#     personal_info.rows[1].cells[1].text = u.place +','+ u.pin +','+ u.district
#     personal_info.rows[2].cells[0].text = 'Phone:'
#     personal_info.rows[2].cells[1].text = u.phone
#
#     # Add a section for educational qualifications
#     document.add_heading('Educational Qualifications', level=1)
#     add_qualification(document, u.qualif)
#     # Add a section for work experience
#     document.add_heading('Work Experience', level=1)
#     add_experience(document, '3')
#
#     l = ', '.join(i.SKILL.skill for i in s)
#
#     # Add a section for skills
#     document.add_heading('Skills', level=1)
#     add_skills(document, l)
#
#     # Save the document
#     date = '/static/cv/' + datetime.datetime.now().strftime('%Y%m%d-%H%M%S') + u.name + '.docx'
#     document.save('C:/Users/Lenovo/PycharmProjects/career_path_navigator/myapp' + date)
#
#     da = datetime.datetime.now().date()
#     if not JobRequest.objects.filter(USER=u, VACCANCY_id=vid).exists():
#         j = JobRequest()
#         j.date = da
#         j.USER = User.objects.get(LOGIN=id)
#         j.VACCANCY_id = vid
#         j.status = 'pending'
#         j.file = date
#         j.save()
#         return JsonResponse({"status": "ok"})
#     return JsonResponse({"status": "no"})
#
# def add_qualification(document, degree):
#     table = document.add_table(rows=1, cols=3)
#     table.alignment = WD_ALIGN_VERTICAL.CENTER
#
#     table.rows[0].cells[0].text = degree
#
# def add_experience(document, title):
#     document.add_paragraph(style='List Bullet').add_run(title).bold = True
#     # document.add_paragraph(description)
#
# def add_skills(document, *skills):
#     skills_str = ', '.join(skills)
#     document.add_paragraph(skills_str)
#
import datetime
import os
from django.http import JsonResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from myapp.models import User, OwnSkill, JobRequest


def create_resume2(request):
    try:
        # Get user details
        id = request.POST.get('lid')
        vid = request.POST.get('vid')

        u = User.objects.get(LOGIN=id)
        s = OwnSkill.objects.filter(USER__LOGIN_id=id)

        document = Document()
        document.add_heading('Resume', 0)

        # Add photo
        photo_path = os.path.join('C:/Users/Lenovo/PycharmProjects/career_path_navigator/', u.photo)
        if os.path.exists(photo_path):
            document.add_picture(photo_path, width=Inches(2.0), height=Inches(2.0))

        # Personal Information
        personal_info = document.add_table(rows=3, cols=2)
        personal_info.alignment = WD_ALIGN_VERTICAL.CENTER
        personal_info.rows[0].cells[0].text = 'Name:'
        personal_info.rows[0].cells[1].text = u.name
        personal_info.rows[1].cells[0].text = 'Address:'
        personal_info.rows[1].cells[1].text = f"{u.place}, {u.pin}, {u.district}"
        personal_info.rows[2].cells[0].text = 'Phone:'
        personal_info.rows[2].cells[1].text = u.phone

        # Education
        document.add_heading('Educational Qualifications', level=1)
        add_qualification(document, str(u.qualif))

        # Work Experience
        document.add_heading('Work Experience', level=1)
        add_experience(document, '3')

        # Skills
        l = ', '.join(str(i.SKILL.skill) for i in s)
        document.add_heading('Skills', level=1)
        add_skills(document, l)

        # Save the document
        timestamp = datetime.datetime.now().strftime('%Y%m%d-%H%M%S')
        filename = f"{timestamp}_{u.name}.docx"
        save_path = f'C:/Users/Lenovo/PycharmProjects/career_path_navigator/myapp/static/cv/{filename}'
        document.save(save_path)

        # Job Request Entry
        da = datetime.datetime.now().date()
        if not JobRequest.objects.filter(USER=u, VACCANCY_id=vid).exists():
            j = JobRequest(
                date=da,
                USER=u,
                VACANCY_id=vid,
                status='pending',
                file=f'/static/cv/{filename}'
            )
            j.save()
            return JsonResponse({"status": "ok"})

        return JsonResponse({"status": "no"})

    except User.DoesNotExist:
        return JsonResponse({"status": "error", "message": "User not found"})
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)})


def add_qualification(document, degree):
    """ Adds education details to the document. """
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[0].cells[0].text = degree


def add_experience(document, title):
    """ Adds work experience details to the document. """
    document.add_paragraph(style='List Bullet').add_run(title).bold = True


def add_skills(document, skills):
    """ Adds skills to the document. """
    document.add_paragraph(skills)
